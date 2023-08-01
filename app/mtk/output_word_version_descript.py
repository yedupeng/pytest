import os
import re
import time

import matplotlib
import matplotlib.font_manager
from docx import Document
from docx.shared import Cm
from git.repo import Repo
from matplotlib import pyplot as plt


def replace_word_text_list(word, temp_text_list, real_text_list):
    for paragraph in word.paragraphs:
        for paragraph_run in paragraph.runs:
            for temp_text in temp_text_list:
                if temp_text in paragraph_run.text:
                    paragraph_run.text = paragraph_run.text.replace(temp_text,
                                                                    real_text_list[temp_text_list.index(temp_text)])

    for table in word.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for paragraph_run in paragraph.runs:
                        for temp_text in temp_text_list:
                            if temp_text in paragraph_run.text:
                                paragraph_run.text = paragraph_run.text.replace(temp_text, real_text_list[
                                    temp_text_list.index(temp_text)])


def replace_word_chart(word, log_list):
    chart_width = word.sections[0].page_width.cm
    for paragraph in word.paragraphs:
        for paragraph_run in paragraph.runs:
            if 'git_chart' in paragraph_run.text:
                chart_str_list = paragraph_run.text.replace('>', '').split(' ')
                paragraph_run.text = ''
                chart_config = {}
                for chart_str in chart_str_list:
                    if '=' in chart_str:
                        key_value = chart_str.split('=')
                        chart_config[key_value[0]] = key_value[1]

                generate_git_chart(log_list, 'temp.png', chart_config['data'], chart_config['type'],
                                   chart_config['title'])
                if '条形图' in chart_config['type']:
                    paragraph_run.add_picture('temp.png', width=Cm(chart_width * 0.7))
                elif '饼图' in chart_config['type']:
                    paragraph_run.add_picture('temp.png', width=Cm(chart_width * 0.6))
    os.remove('temp.png')


def get_git_key_str(summary, key):
    key_list = ['网关', '运营商', '省份', '产品型号', '类型', '进程', '单号', '描述', '测试', '开发者', '检视人']
    key_obj = re.findall(r'\[.*?]', summary)
    if len(key_obj) < len(key_list):
        return ''
    if key_list.index(key) < key_list.index('单号'):
        return key_obj[key_list.index(key)].replace('[', '').replace(']', '').replace(' ', '').upper()
    else:
        for key_temp in key_obj:
            if f'{key}：' in key_temp:
                return key_temp.split('：', 1)[1].replace(']', '').replace(' ', '').upper()
        return ''


def generate_git_statistic(log_list, key):
    key_list = []
    value_list = []
    for item in log_list:
        key_str_list = get_git_key_str(item.get("summary"), key).split('|')
        if len(key_str_list) <= 0:
            continue
        for key_str in key_str_list:
            if len(key_str) <= 0:
                continue
            if key_str not in key_list:
                value_list.append(1)
                key_list.append(key_str)
            else:
                value_list[key_list.index(key_str)] += 1
    value_list, key_list = (list(item) for item in zip(*sorted(zip(value_list, key_list), reverse=True)))
    return key_list, value_list


def generate_git_chart(log_list, picture_path, key, picture_style, picture_title):
    git_key_list, git_value_list = generate_git_statistic(log_list, key)
    color_list = ['#95a2ff', '#fa8080', '#ffc076', '#fae768', '#87e885', '#3cb9fc', '#73abf5', '#cb9bff',
                  '#90ed7d', '#f7a35c', '#8085e9']
    if picture_style == '条形图':
        plt.figure(figsize=(10, 6))
        if len(git_key_list) <= 8:
            pass
        elif len(git_key_list) <= 16:
            plt.xticks(rotation=30)
        elif len(git_key_list) <= 25:
            plt.xticks(rotation=60)
        else:
            plt.xticks(rotation=90)
        bar = plt.bar(git_key_list, git_value_list, width=0.4, color=color_list)
        plt.title(picture_title)
        plt.bar_label(bar, label_type='edge')
        plt.xticks()
        plt.yticks()
        plt.savefig(picture_path, dpi=300, bbox_inches='tight')
    elif picture_style == '饼图':
        plt.figure(figsize=(8, 5))
        plt.pie(git_value_list, labels=git_key_list, autopct='%d%%', colors=color_list)
        plt.title(picture_title)
        plt.savefig(picture_path, dpi=300, bbox_inches='tight')


def run(project_config, option_config):
    matplotlib.font_manager.fontManager.addfont('etc/simhei.ttf')
    plt.rcParams['axes.unicode_minus'] = False
    plt.rcParams['font.sans-serif'] = ['SimHei']
    word_temp_path = option_config['temp_docx']
    word_path = f'{project_config["pack_path"]}/{project_config["product"]}_{project_config["chip"]}' \
                f'_V{project_config["version"]}版本描述文件.docx'

    log_str_list = Repo(project_config['path']).git.log(
        "--pretty={'''commit''':'''%H''','''author''':'''%an''','''summary''':'''%s''','''date''':'''%cd'''}",
        date='format:%Y-%m-%d %H:%M').split("\n")
    all_log_list = [eval(item) for item in log_str_list]
    version_log_list = []
    git_log_str = ''
    for item in all_log_list:
        version_log_list.append(item)
        author_name = item.get("author")
        git_log_str = git_log_str + f'{item.get("commit")}  ' \
                                    f'{author_name}  {item.get("date")}\n    {item.get("summary")}\n\n\n\n '
        if item.get("commit") == option_config['first_version_commit_id'].strip():
            break

    extra_temp_dict = {'date': time.strftime("%Y-%m-%d", time.localtime()),
                       'git_id': Repo(project_config['path']).git.log('--pretty=%H', max_count=1),
                       'git_log': git_log_str,
                       'device_version': project_config['device_version_str'],
                       'device_info': project_config['device_info_str']}

    temp_text_list = []
    real_text_list = []
    for temp_word in option_config['replace_word']:
        temp_text_list.append(f'<{temp_word}>')
        real_text_list.append(option_config['replace_word'][temp_word])

    for temp_word in project_config:
        temp_text_list.append(f'<{temp_word}>')
        real_text_list.append(project_config[temp_word])

    for temp_word in extra_temp_dict:
        temp_text_list.append(f'<{temp_word}>')
        real_text_list.append(extra_temp_dict[temp_word])

    word = Document(word_temp_path)
    replace_word_text_list(word, temp_text_list, real_text_list)
    replace_word_chart(word, version_log_list)
    word.save(word_path)

    return 0
