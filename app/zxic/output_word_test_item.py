import os
from docx import Document


def replace_word_text_list(word, temp_text_list, real_text_list):
    for paragraph in word.paragraphs:
        for paragraph_run in paragraph.runs:
            for temp_text in temp_text_list:
                if temp_text in paragraph_run.text:
                    paragraph_run.text = paragraph_run.text.replace(temp_text,
                                                                    real_text_list[temp_text_list.index(temp_text)])

    for table in word.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for paragraph_run in paragraph.runs:
                        for temp_text in temp_text_list:
                            if temp_text in paragraph_run.text:
                                paragraph_run.text = paragraph_run.text.replace(temp_text, real_text_list[
                                    temp_text_list.index(temp_text)])


def run(project_config, option_config):
    word_temp_path = option_config['temp_docx']
    word_path = f'{project_config["pack_path"]}/{project_config["product"]}_{project_config["chip"]}' \
                f'_V{project_config["version"]}提测事项表.docx'

    extra_temp_dict = {}

    for device_info_dir in os.listdir(f'{project_config["path"]}/Release/{project_config["pon"]}'):
        if 'autoMake' in device_info_dir:
            for line in open(
                    f'{project_config["path"]}/Release/{project_config["pon"]}/{device_info_dir}/original_device_info',
                    'r'):
                key_value = line.split('=')
                extra_temp_dict[f'{key_value[0]}'] = key_value[1].replace('\n', '')

    soft_version_path = f'{project_config["path"]}/compatible_branch/make/config_cmcc/' \
                        f'{project_config["pon"]}/softversion_config'
    if os.path.exists(soft_version_path):
        for soft_version_list in open(soft_version_path, 'r'):
            key_value = soft_version_list.split('=')
            extra_temp_dict[f'{key_value[0]}_soft_version'] = key_value[1].replace('.cmp', '').replace('\n', '')

    temp_text_list = []
    real_text_list = []
    for temp_word in option_config['replace_word']:
        temp_text_list.append(f'<{temp_word}>')
        real_text_list.append(option_config['replace_word'][temp_word])

    for temp_word in project_config:
        temp_text_list.append(f'<{temp_word}>')
        real_text_list.append(project_config[temp_word])

    for temp_word in extra_temp_dict:
        temp_text_list.append(f'<{temp_word}>')
        real_text_list.append(extra_temp_dict[temp_word])

    word = Document(word_temp_path)
    replace_word_text_list(word, temp_text_list, real_text_list)
    word.save(word_path)

    return 0
